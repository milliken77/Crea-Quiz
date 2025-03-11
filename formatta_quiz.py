import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# Funzione per aggiungere una tabella di correzione
def add_griglia_correzione(doc, num_domande):
    # Aggiungi riga vuota prima della griglia di correzione senza interlinea
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    doc.add_paragraph("Griglia di correzione").bold = True
    table = doc.add_table(rows=num_domande + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Imposta larghezza delle colonne e altezza delle righe (altezza ridotta del 30%)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(2.5)
        row.height = Cm(0.6)  # Altezza ridotta del 30%

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Domande'
    hdr_cells[1].text = 'Risposte'
    
    for i in range(1, num_domande + 1):
        table.cell(i, 0).text = str(i)
        table.cell(i, 1).text = ''  # Lasciato vuoto per essere compilato

    # Centra il contenuto delle celle
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Funzione per trasformare ^numero in esponente
def format_superscript(paragraph, text):
    parts = text.split('^')
    run = paragraph.add_run(parts[0])  # Prima parte senza esponenti
    for part in parts[1:]:
        if part and part[0].isdigit():  # Se inizia con un numero
            sup_run = paragraph.add_run(part[0])  # Crea il run per l'esponente
            sup_run.font.superscript = True  # Imposta l'apice (esponente)
            paragraph.add_run(part[1:])  # Continua il testo normale
        else:
            paragraph.add_run('^' + part)  # Caso speciale se ^ è seguito da non numeri

# Funzione per aggiungere le domande e risposte senza righe vuote e interlinea
def add_domande_e_risposte(doc, compito):
    domande = compito.strip().split('\n')
    for line in domande:
        if line.strip():
            # Aggiungi riga vuota prima delle domande che iniziano con numero progressivo senza interlinea
            if line.lstrip().startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9', '0')):
                empty_paragraph = doc.add_paragraph()
                empty_paragraph.paragraph_format.space_before = Pt(0)
                empty_paragraph.paragraph_format.space_after = Pt(0)
                empty_paragraph.paragraph_format.line_spacing = 1

            paragraph = doc.add_paragraph()  # Aggiungi un nuovo paragrafo
            
            # Verifica se la linea inizia con un numero e metti in grassetto
            if line.lstrip().startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9', '0')):
                run = paragraph.add_run()
                format_superscript(paragraph, line)  # Formatta con esponenti
                paragraph.runs[0].bold = True  # Grassetto per le domande numerate
            else:
                format_superscript(paragraph, line)  # Formatta con esponenti per le altre righe
            
            # Imposta interlinea compatta
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

# Funzione per creare il documento
def create_compito_word(file_path, text_data):
    doc = Document()
    
    # Imposta due colonne
    section = doc.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Due colonne

    # Funzione per aggiungere il titolo e il form
    def add_header_and_form(doc):
        title = doc.add_paragraph("Compito ###")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(16)  # Font leggermente più grande

        for label in ["Nome:", "Cognome:", "Data:"]:
            paragraph = doc.add_paragraph(f"{label} _______________________")
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

        # Aggiungi una riga vuota senza interlinea tra "Data" e "Compito n°"
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1

    # Split del testo su base ::::::
    compiti = text_data.split(":::::")
    
    for idx, compito in enumerate(compiti):
        if compito.strip():  # Evita le stringhe vuote
            # Aggiungi riga vuota prima della dicitura "Compito n°" senza interlinea
            empty_paragraph = doc.add_paragraph()
            empty_paragraph.paragraph_format.space_before = Pt(0)
            empty_paragraph.paragraph_format.space_after = Pt(0)
            empty_paragraph.paragraph_format.line_spacing = 1

            add_header_and_form(doc)
            
            # Aggiungi domande senza righe vuote tra di esse
            add_domande_e_risposte(doc, compito)

            num_domande = compito.strip().count("A : ")  # Conta le domande
            add_griglia_correzione(doc, num_domande)
            if idx < len(compiti) - 1:
                doc.add_page_break()  # Aggiunge nuova pagina se non è l'ultimo compito
    
    # Salva il documento
    doc.save(file_path)

# Funzione principale per leggere il file di testo e generare il documento
def main(input_txt, output_docx):
    with open(input_txt, 'r', encoding='utf-8') as file:
        text_data = file.read()
    
    create_compito_word(output_docx, text_data)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Crea un file Word formattato da un file di testo con compiti.")
    parser.add_argument("input_txt", help="Il percorso del file di testo contenente i compiti.")
    parser.add_argument("output_docx", help="Il percorso dove salvare il file Word generato.")
    
    args = parser.parse_args()
    main(args.input_txt, args.output_docx)

